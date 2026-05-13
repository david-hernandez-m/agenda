import json
import os
import re
import webbrowser
from calendar import month_name, monthcalendar
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk

APP_NAME = "Mi Agenda Personal Premium"
BASE_DIR = Path.home() / "MiAgendaPersonal"
ENTRIES_DIR = BASE_DIR / "entradas"
EXPORTS_DIR = BASE_DIR / "exportados"
CONFIG_FILE = BASE_DIR / "config.json"
AUTOSTART_BAT = BASE_DIR / "iniciar_agenda.bat"

try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.pdfgen import canvas
    PDF_OK = True
except Exception:
    PDF_OK = False


DEFAULT_CONFIG = {
    "cloud_folder": "",
    "last_window_x": None,
    "last_window_y": None,
    "theme": "dark",
    "pin_enabled": False,
    "pin_value": "",
    "autostart_enabled": False
}


def ensure_dirs():
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    ENTRIES_DIR.mkdir(parents=True, exist_ok=True)
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    if not CONFIG_FILE.exists():
        CONFIG_FILE.write_text(
            json.dumps(DEFAULT_CONFIG, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )


def load_config():
    ensure_dirs()
    try:
        data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        merged = DEFAULT_CONFIG.copy()
        merged.update(data)
        return merged
    except Exception:
        return DEFAULT_CONFIG.copy()


def save_config(data):
    CONFIG_FILE.write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )


def slugify(text):
    text = text.strip().lower()
    text = re.sub(r"[^a-zA-Z0-9áéíóúñÁÉÍÓÚÑ _-]", "", text)
    text = text.replace(" ", "_")
    text = re.sub(r"_+", "_", text)
    return text[:50] if text else "entrada"


def current_stamp():
    return datetime.now().strftime("%Y-%m-%d_%H-%M-%S")


def wrap_text_pdf(text, font_name, font_size, max_width):
    words = text.split()
    lines = []
    current = ""

    for word in words:
        probe = word if not current else current + " " + word
        if stringWidth(probe, font_name, font_size) <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines if lines else [""]


def save_txt(path, content):
    path.write_text(content, encoding="utf-8")


def save_pdf(path, content):
    if not PDF_OK:
        raise RuntimeError("Falta instalar reportlab. Usa: pip install reportlab")

    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 2 * cm
    y = height - margin
    usable_width = width - (2 * margin)
    line_height = 16
    font_name = "Helvetica"
    font_size = 11

    c.setTitle(APP_NAME)

    for paragraph in content.splitlines():
        wrapped = wrap_text_pdf(paragraph, font_name, font_size, usable_width) if paragraph else [""]

        for line in wrapped:
            if y < margin:
                c.showPage()
                y = height - margin
            c.setFont(font_name, font_size)
            c.drawString(margin, y, line)
            y -= line_height

        if not paragraph:
            y -= 4

    c.save()


def save_docx(path, title, feeling, body, created_at):
    if not DOCX_OK:
        raise RuntimeError("Falta instalar python-docx. Usa: pip install python-docx")

    doc = Document()
    doc.add_heading(APP_NAME, 0)
    doc.add_paragraph(f"Fecha: {created_at}")
    doc.add_paragraph(f"Estado del día: {feeling}")
    doc.add_paragraph(f"Título: {title}")
    doc.add_paragraph("")
    doc.add_paragraph(body)
    doc.save(str(path))


class AgendaApp:
    def __init__(self, root):
        self.root = root
        self.config = load_config()
        self.current_entry_file = None
        self.history_files = []

        self.feeling_var = tk.StringVar(value="Bien")
        self.title_var = tk.StringVar()
        self.word_count_var = tk.StringVar(value="Palabras: 0")
        self.char_count_var = tk.StringVar(value="Caracteres: 0")
        self.status_var = tk.StringVar(value="Listo")
        self.calendar_month_var = tk.IntVar(value=datetime.now().month)
        self.calendar_year_var = tk.IntVar(value=datetime.now().year)

        self.root.title(APP_NAME)
        self.root.geometry("1280x760")
        self.root.minsize(1080, 680)

        self.apply_window_position()
        self.setup_theme()
        self.build_ui()
        self.refresh_history()
        self.render_calendar()

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        if self.config.get("pin_enabled"):
            self.ask_pin_on_start()

    # =========================
    # THEME
    # =========================
    def setup_theme(self):
        self.style = ttk.Style()
        if "clam" in self.style.theme_names():
            self.style.theme_use("clam")
        self.apply_theme()

    def apply_theme(self):
        dark = self.config.get("theme", "dark") == "dark"

        if dark:
            colors = {
                "bg": "#141821",
                "panel": "#1c2230",
                "panel2": "#20293a",
                "text": "#f3f6ff",
                "subtext": "#aeb8cf",
                "entry_bg": "#0f1420",
                "accent": "#4f8cff",
                "sel": "#2a4678"
            }
        else:
            colors = {
                "bg": "#eef2f8",
                "panel": "#ffffff",
                "panel2": "#e9eef8",
                "text": "#1d2433",
                "subtext": "#5e6678",
                "entry_bg": "#ffffff",
                "accent": "#3a73e8",
                "sel": "#cfe0ff"
            }

        self.colors = colors
        self.root.configure(bg=colors["bg"])

        self.style.configure("Main.TFrame", background=colors["bg"])
        self.style.configure("Panel.TFrame", background=colors["panel"])
        self.style.configure("Side.TFrame", background=colors["panel2"])
        self.style.configure("Title.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 20, "bold"))
        self.style.configure("Sub.TLabel", background=colors["panel"], foreground=colors["subtext"], font=("Segoe UI", 10))
        self.style.configure("Panel.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10))
        self.style.configure("SideTitle.TLabel", background=colors["panel2"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        self.style.configure("Accent.TButton", font=("Segoe UI", 10, "bold"))
        self.style.map("Accent.TButton", background=[("active", colors["accent"])])

    def refresh_theme_widgets(self):
        dark = self.config.get("theme", "dark") == "dark"
        fg = self.colors["text"]
        bg = self.colors["entry_bg"]
        insert_bg = "#ffffff" if dark else "#000000"

        self.text.configure(
            bg=bg,
            fg=fg,
            insertbackground=insert_bg,
            selectbackground=self.colors["sel"]
        )
        self.history_list.configure(
            bg=bg,
            fg=fg,
            selectbackground=self.colors["sel"],
            selectforeground=fg
        )
        self.calendar_text.configure(
            bg=bg,
            fg=fg,
            insertbackground=insert_bg
        )

    # =========================
    # WINDOW
    # =========================
    def apply_window_position(self):
        x = self.config.get("last_window_x")
        y = self.config.get("last_window_y")
        if isinstance(x, int) and isinstance(y, int):
            self.root.geometry(f"1280x760+{x}+{y}")
        else:
            self.center_window(1280, 760)

    def center_window(self, w, h):
        self.root.update_idletasks()
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        x = int((sw / 2) - (w / 2))
        y = int((sh / 2) - (h / 2))
        self.root.geometry(f"{w}x{h}+{x}+{y}")

    # =========================
    # UI
    # =========================
    def build_ui(self):
        main = ttk.Frame(self.root, style="Main.TFrame", padding=10)
        main.pack(fill="both", expand=True)

        layout = ttk.Panedwindow(main, orient="horizontal")
        layout.pack(fill="both", expand=True)

        left = ttk.Frame(layout, style="Side.TFrame", padding=12)
        center = ttk.Frame(layout, style="Panel.TFrame", padding=16)
        right = ttk.Frame(layout, style="Side.TFrame", padding=12)

        layout.add(left, weight=1)
        layout.add(center, weight=4)
        layout.add(right, weight=1)

        # LEFT
        ttk.Label(left, text="Historial", style="SideTitle.TLabel").pack(anchor="w", pady=(0, 8))

        ttk.Button(left, text="Nueva entrada", command=self.new_entry).pack(fill="x", pady=2)
        ttk.Button(left, text="Guardar entrada", command=self.save_entry).pack(fill="x", pady=2)
        ttk.Button(left, text="Guardar y limpiar", command=self.save_and_clear).pack(fill="x", pady=2)
        ttk.Button(left, text="Abrir carpeta", command=self.open_entries_folder).pack(fill="x", pady=2)

        self.history_list = tk.Listbox(
            left,
            font=("Segoe UI", 10),
            relief="flat",
            highlightthickness=0,
            activestyle="none"
        )
        self.history_list.pack(fill="both", expand=True, pady=(10, 0))
        self.history_list.bind("<<ListboxSelect>>", self.on_history_select)

        # CENTER
        center.columnconfigure(0, weight=1)
        center.rowconfigure(4, weight=1)

        ttk.Label(center, text="¿Cómo ha estado tu día hoy?", style="Title.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(center, text="Escribe libremente lo que quieras guardar, exportar o compartir.", style="Sub.TLabel").grid(
            row=1, column=0, sticky="w", pady=(2, 12)
        )

        top = ttk.Frame(center, style="Panel.TFrame")
        top.grid(row=2, column=0, sticky="ew", pady=(0, 10))
        top.columnconfigure(3, weight=1)

        ttk.Label(top, text="Estado del día:", style="Panel.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Combobox(
            top,
            textvariable=self.feeling_var,
            state="readonly",
            width=18,
            values=["Muy bien", "Bien", "Normal", "Cansado/a", "Triste", "Ansioso/a", "Enojado/a", "Otro"]
        ).grid(row=0, column=1, sticky="w", padx=(8, 16))

        ttk.Label(top, text="Título:", style="Panel.TLabel").grid(row=0, column=2, sticky="w")
        ttk.Entry(top, textvariable=self.title_var).grid(row=0, column=3, sticky="ew", padx=(8, 0))

        ttk.Label(center, text="Tu texto:", style="Panel.TLabel").grid(row=3, column=0, sticky="w")

        text_wrap = ttk.Frame(center, style="Panel.TFrame")
        text_wrap.grid(row=4, column=0, sticky="nsew", pady=(6, 8))
        text_wrap.rowconfigure(0, weight=1)
        text_wrap.columnconfigure(0, weight=1)

        self.text = tk.Text(
            text_wrap,
            wrap="word",
            font=("Segoe UI", 11),
            undo=True,
            relief="flat",
            bd=0
        )
        self.text.grid(row=0, column=0, sticky="nsew")

        text_scroll = ttk.Scrollbar(text_wrap, orient="vertical", command=self.text.yview)
        text_scroll.grid(row=0, column=1, sticky="ns")
        self.text.configure(yscrollcommand=text_scroll.set)
        self.text.bind("<KeyRelease>", self.on_text_changed)

        info = ttk.Frame(center, style="Panel.TFrame")
        info.grid(row=5, column=0, sticky="ew", pady=(0, 8))

        ttk.Label(info, textvariable=self.word_count_var, style="Panel.TLabel").pack(side="left")
        ttk.Label(info, textvariable=self.char_count_var, style="Panel.TLabel").pack(side="left", padx=(14, 0))

        actions = ttk.LabelFrame(center, text="Acciones")
        actions.grid(row=6, column=0, sticky="ew", pady=(6, 0))
        for i in range(4):
            actions.columnconfigure(i, weight=1)

        ttk.Button(actions, text="Enviar / Exportar", style="Accent.TButton", command=self.show_export_menu).grid(
            row=0, column=0, sticky="ew", padx=6, pady=10
        )
        ttk.Button(actions, text="WhatsApp", style="Accent.TButton", command=self.send_whatsapp).grid(
            row=0, column=1, sticky="ew", padx=6, pady=10
        )
        ttk.Button(actions, text="Correo", command=self.send_email).grid(
            row=0, column=2, sticky="ew", padx=6, pady=10
        )
        ttk.Button(actions, text="Borrar formulario", command=self.clear_form).grid(
            row=0, column=3, sticky="ew", padx=6, pady=10
        )

        status = ttk.Frame(center, style="Panel.TFrame")
        status.grid(row=7, column=0, sticky="ew", pady=(10, 0))
        ttk.Label(status, textvariable=self.status_var, style="Sub.TLabel").pack(side="left")

        # RIGHT
        ttk.Label(right, text="Herramientas", style="SideTitle.TLabel").pack(anchor="w", pady=(0, 8))

        ttk.Button(right, text="Cambiar tema", command=self.toggle_theme).pack(fill="x", pady=2)
        ttk.Button(right, text="Configurar nube", command=self.set_cloud_folder).pack(fill="x", pady=2)
        ttk.Button(right, text="PIN de acceso", command=self.configure_pin).pack(fill="x", pady=2)
        ttk.Button(right, text="Inicio con Windows", command=self.toggle_autostart).pack(fill="x", pady=2)

        ttk.Label(right, text="Calendario", style="SideTitle.TLabel").pack(anchor="w", pady=(14, 6))

        nav = ttk.Frame(right, style="Side.TFrame")
        nav.pack(fill="x", pady=(0, 6))

        ttk.Button(nav, text="<", width=3, command=self.prev_month).pack(side="left")
        ttk.Button(nav, text=">", width=3, command=self.next_month).pack(side="right")

        self.calendar_header = ttk.Label(right, text="", style="Panel.TLabel")
        self.calendar_header.pack(anchor="center", pady=(0, 4))

        self.calendar_text = tk.Text(
            right,
            height=10,
            width=22,
            wrap="none",
            font=("Consolas", 10),
            relief="flat",
            bd=0
        )
        self.calendar_text.pack(fill="x", pady=(0, 10))
        self.calendar_text.configure(state="disabled")

        ttk.Label(right, text="Exportar rápido", style="SideTitle.TLabel").pack(anchor="w", pady=(6, 6))
        ttk.Button(right, text="Guardar TXT", command=lambda: self.save_as("txt")).pack(fill="x", pady=2)
        ttk.Button(right, text="Guardar PDF", command=lambda: self.save_as("pdf")).pack(fill="x", pady=2)
        ttk.Button(right, text="Guardar Word", command=lambda: self.save_as("docx")).pack(fill="x", pady=2)

        self.refresh_theme_widgets()

    # =========================
    # PIN
    # =========================
    def ask_pin_on_start(self):
        pin = simpledialog.askstring("PIN", "Ingresa tu PIN:", show="*")
        if pin != self.config.get("pin_value", ""):
            messagebox.showerror("Acceso denegado", "PIN incorrecto.")
            self.root.after(100, self.root.destroy)

    def configure_pin(self):
        if self.config.get("pin_enabled"):
            choice = messagebox.askyesnocancel("PIN", "¿Quieres desactivar el PIN?\n\nSí = desactivar\nNo = cambiar PIN")
            if choice is None:
                return
            if choice:
                self.config["pin_enabled"] = False
                self.config["pin_value"] = ""
                save_config(self.config)
                self.status_var.set("PIN desactivado")
                messagebox.showinfo("PIN", "PIN desactivado.")
                return

        new_pin = simpledialog.askstring("Configurar PIN", "Escribe un PIN numérico:", show="*")
        if not new_pin:
            return
        if not new_pin.isdigit() or len(new_pin) < 4:
            messagebox.showwarning("PIN inválido", "Usa un PIN numérico de al menos 4 dígitos.")
            return

        confirm = simpledialog.askstring("Confirmar PIN", "Repite el PIN:", show="*")
        if new_pin != confirm:
            messagebox.showerror("Error", "Los PIN no coinciden.")
            return

        self.config["pin_enabled"] = True
        self.config["pin_value"] = new_pin
        save_config(self.config)
        self.status_var.set("PIN configurado")
        messagebox.showinfo("PIN", "PIN guardado correctamente.")

    # =========================
    # CALENDAR
    # =========================
    def render_calendar(self):
        month = self.calendar_month_var.get()
        year = self.calendar_year_var.get()
        self.calendar_header.config(text=f"{month_name[month]} {year}")

        lines = ["Lu Ma Mi Ju Vi Sá Do"]
        cal = monthcalendar(year, month)
        today = datetime.now()

        for week in cal:
            parts = []
            for day in week:
                if day == 0:
                    parts.append("  ")
                else:
                    mark = f"{day:2d}"
                    if day == today.day and month == today.month and year == today.year:
                        mark = f"[{day:02d}]"
                    parts.append(mark)
            line = " ".join(parts).replace("[", "").replace("]", "")
            lines.append(line)

        self.calendar_text.configure(state="normal")
        self.calendar_text.delete("1.0", "end")
        self.calendar_text.insert("1.0", "\n".join(lines))
        self.calendar_text.configure(state="disabled")

    def prev_month(self):
        m = self.calendar_month_var.get()
        y = self.calendar_year_var.get()
        if m == 1:
            self.calendar_month_var.set(12)
            self.calendar_year_var.set(y - 1)
        else:
            self.calendar_month_var.set(m - 1)
        self.render_calendar()

    def next_month(self):
        m = self.calendar_month_var.get()
        y = self.calendar_year_var.get()
        if m == 12:
            self.calendar_month_var.set(1)
            self.calendar_year_var.set(y + 1)
        else:
            self.calendar_month_var.set(m + 1)
        self.render_calendar()

    # =========================
    # ENTRY DATA
    # =========================
    def get_body(self):
        return self.text.get("1.0", "end").strip()

    def build_entry_payload(self):
        return {
            "created_at": datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
            "feeling": self.feeling_var.get().strip() or "Sin especificar",
            "title": self.title_var.get().strip() or "Sin título",
            "body": self.get_body()
        }

    def build_export_content(self, payload):
        return (
            f"{APP_NAME}\n"
            f"{'=' * len(APP_NAME)}\n"
            f"Fecha: {payload['created_at']}\n"
            f"Estado del día: {payload['feeling']}\n"
            f"Título: {payload['title']}\n\n"
            f"Escrito:\n{payload['body']}\n"
        )

    def validate_content(self):
        if not self.get_body():
            messagebox.showwarning("Falta contenido", "Escribe algo antes de guardar o enviar.")
            return False
        return True

    def make_json_filename(self, title):
        return f"{current_stamp()}_{slugify(title)}.json"

    # =========================
    # HISTORY
    # =========================
    def refresh_history(self):
        self.history_list.delete(0, "end")
        self.history_files = sorted(ENTRIES_DIR.glob("*.json"), reverse=True)

        for file in self.history_files:
            try:
                data = json.loads(file.read_text(encoding="utf-8"))
                date = data.get("created_at", "")
                title = data.get("title", "Sin título")
                self.history_list.insert("end", f"{date} | {title}")
            except Exception:
                self.history_list.insert("end", file.name)

    def on_history_select(self, event=None):
        sel = self.history_list.curselection()
        if not sel:
            return

        idx = sel[0]
        file = self.history_files[idx]
        try:
            data = json.loads(file.read_text(encoding="utf-8"))
            self.feeling_var.set(data.get("feeling", "Bien"))
            self.title_var.set(data.get("title", ""))
            self.text.delete("1.0", "end")
            self.text.insert("1.0", data.get("body", ""))
            self.current_entry_file = file
            self.update_counts()
            self.status_var.set(f"Entrada cargada: {file.name}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la entrada.\n\n{e}")

    # =========================
    # ACTIONS
    # =========================
    def save_entry(self):
        if not self.validate_content():
            return

        payload = self.build_entry_payload()
        if self.current_entry_file:
            path = self.current_entry_file
        else:
            path = ENTRIES_DIR / self.make_json_filename(payload["title"])
            self.current_entry_file = path

        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        self.refresh_history()
        self.status_var.set(f"Entrada guardada: {path.name}")
        messagebox.showinfo("Guardado", "La entrada se guardó correctamente.")

    def save_and_clear(self):
        self.save_entry()
        if self.current_entry_file is not None:
            self.clear_form()
            self.status_var.set("Entrada guardada y formulario limpiado")

    def new_entry(self):
        self.clear_form()
        self.status_var.set("Nueva entrada creada")

    def clear_form(self):
        self.feeling_var.set("Bien")
        self.title_var.set("")
        self.text.delete("1.0", "end")
        self.current_entry_file = None
        self.update_counts()
        self.status_var.set("Formulario limpio")

    def on_text_changed(self, event=None):
        self.update_counts()
        self.status_var.set("Editando...")

    def update_counts(self):
        body = self.get_body()
        words = len(body.split()) if body else 0
        chars = len(body)
        self.word_count_var.set(f"Palabras: {words}")
        self.char_count_var.set(f"Caracteres: {chars}")

    def open_entries_folder(self):
        webbrowser.open(str(ENTRIES_DIR))

    def set_cloud_folder(self):
        folder = filedialog.askdirectory(title="Selecciona tu carpeta sincronizada")
        if folder:
            self.config["cloud_folder"] = folder
            save_config(self.config)
            self.status_var.set("Carpeta de nube configurada")
            messagebox.showinfo("Nube", f"Carpeta guardada:\n{folder}")

    def toggle_theme(self):
        self.config["theme"] = "light" if self.config.get("theme") == "dark" else "dark"
        save_config(self.config)
        self.apply_theme()
        self.refresh_theme_widgets()
        self.status_var.set(f"Tema cambiado a {self.config['theme']}")

    def show_export_menu(self):
        if not self.validate_content():
            return

        win = tk.Toplevel(self.root)
        win.title("Enviar / Exportar")
        win.geometry("420x360")
        win.transient(self.root)
        win.grab_set()

        frm = ttk.Frame(win, padding=14)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="Elige qué deseas hacer", font=("Segoe UI", 13, "bold")).pack(anchor="w", pady=(0, 10))
        ttk.Button(frm, text="Enviar por correo", command=lambda: [win.destroy(), self.send_email()]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Enviar por WhatsApp", command=lambda: [win.destroy(), self.send_whatsapp()]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar en nube", command=lambda: [win.destroy(), self.save_to_cloud_menu()]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar en TXT", command=lambda: [win.destroy(), self.save_as("txt")]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar en PDF", command=lambda: [win.destroy(), self.save_as("pdf")]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar en Word", command=lambda: [win.destroy(), self.save_as("docx")]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Cancelar", command=win.destroy).pack(fill="x", pady=(10, 0))

    def save_to_cloud_menu(self):
        cloud_folder = self.config.get("cloud_folder", "")
        if not cloud_folder or not Path(cloud_folder).exists():
            messagebox.showwarning("Nube no configurada", "Primero configura una carpeta sincronizada.")
            self.set_cloud_folder()
            cloud_folder = self.config.get("cloud_folder", "")
            if not cloud_folder:
                return

        win = tk.Toplevel(self.root)
        win.title("Guardar en nube")
        win.geometry("360x220")
        win.transient(self.root)
        win.grab_set()

        frm = ttk.Frame(win, padding=14)
        frm.pack(fill="both", expand=True)

        ttk.Label(frm, text="Selecciona el formato", font=("Segoe UI", 11, "bold")).pack(anchor="w", pady=(0, 10))
        ttk.Button(frm, text="Guardar TXT", command=lambda: [win.destroy(), self.save_as("txt", Path(cloud_folder))]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar PDF", command=lambda: [win.destroy(), self.save_as("pdf", Path(cloud_folder))]).pack(fill="x", pady=4)
        ttk.Button(frm, text="Guardar Word", command=lambda: [win.destroy(), self.save_as("docx", Path(cloud_folder))]).pack(fill="x", pady=4)

    def save_as(self, fmt, folder=None):
        if not self.validate_content():
            return

        payload = self.build_entry_payload()
        default_name = f"{current_stamp()}_{slugify(payload['title'])}"

        ext_map = {"txt": ".txt", "pdf": ".pdf", "docx": ".docx"}
        ext = ext_map[fmt]

        if folder is None:
            file_path = filedialog.asksaveasfilename(
                title=f"Guardar como {fmt.upper()}",
                defaultextension=ext,
                initialfile=default_name + ext,
                filetypes=[(fmt.upper(), f"*{ext}"), ("Todos los archivos", "*.*")]
            )
            if not file_path:
                return
            path = Path(file_path)
        else:
            folder.mkdir(parents=True, exist_ok=True)
            path = folder / (default_name + ext)

        try:
            if fmt == "txt":
                save_txt(path, self.build_export_content(payload))
            elif fmt == "pdf":
                save_pdf(path, self.build_export_content(payload))
            elif fmt == "docx":
                save_docx(path, payload["title"], payload["feeling"], payload["body"], payload["created_at"])
            else:
                raise ValueError("Formato no soportado")

            self.status_var.set(f"Archivo exportado: {path.name}")
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el archivo.\n\n{e}")

    def send_email(self):
        if not self.validate_content():
            return

        payload = self.build_entry_payload()
        subject = payload["title"]
        body = self.build_export_content(payload)
        url = f"mailto:?subject={quote(subject)}&body={quote(body)}"
        webbrowser.open(url)
        self.status_var.set("Correo preparado")

    def send_whatsapp(self):
        if not self.validate_content():
            return

        payload = self.build_entry_payload()
        body = self.build_export_content(payload)

        answer = messagebox.askyesno(
            "WhatsApp",
            "¿Quieres escribir un número?\n\nSí = escribir número\nNo = abrir WhatsApp Web con el mensaje general"
        )

        if answer:
            number = simpledialog.askstring(
                "Número de WhatsApp",
                "Escribe el número con código país.\nEjemplo Chile: 56912345678"
            )
            if not number:
                return
            number = re.sub(r"\D", "", number)
            url = f"https://wa.me/{number}?text={quote(body)}"
        else:
            url = f"https://web.whatsapp.com/send?text={quote(body)}"

        webbrowser.open(url)
        self.status_var.set("WhatsApp Web abierto")

    # =========================
    # AUTOSTART WINDOWS
    # =========================
    def toggle_autostart(self):
        startup_dir = Path(os.getenv("APPDATA", "")) / "Microsoft" / "Windows" / "Start Menu" / "Programs" / "Startup"
        startup_bat = startup_dir / "MiAgendaPersonal.bat"

        if startup_bat.exists():
            try:
                startup_bat.unlink()
                self.config["autostart_enabled"] = False
                save_config(self.config)
                self.status_var.set("Inicio automático desactivado")
                messagebox.showinfo("Inicio con Windows", "Se desactivó el inicio automático.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo desactivar.\n\n{e}")
            return

        try:
            python_exe = Path(os.sys.executable)
            script_path = Path(__file__).resolve()
            content = f'@echo off\n"{python_exe}" "{script_path}"\n'
            startup_bat.write_text(content, encoding="utf-8")
            self.config["autostart_enabled"] = True
            save_config(self.config)
            self.status_var.set("Inicio automático activado")
            messagebox.showinfo("Inicio con Windows", "La app se abrirá al iniciar sesión en Windows.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo activar el inicio automático.\n\n{e}")

    # =========================
    # CLOSE
    # =========================
    def on_close(self):
        try:
            self.config["last_window_x"] = self.root.winfo_x()
            self.config["last_window_y"] = self.root.winfo_y()
            save_config(self.config)
        finally:
            self.root.destroy()


def main():
    ensure_dirs()
    root = tk.Tk()
    app = AgendaApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()